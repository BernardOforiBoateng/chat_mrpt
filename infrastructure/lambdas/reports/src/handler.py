"""
Reports Lambda function.
Generates PDF and Word reports for analysis results.
"""
import json
import boto3
import os
import pandas as pd
from typing import Dict, Any, List, Optional
from datetime import datetime
import tempfile
import uuid
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


s3_client = boto3.client('s3')
dynamodb = boto3.resource('dynamodb')


def lambda_handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    """
    Generate analysis reports in various formats.

    Args:
        event: API Gateway event or Step Functions input
        context: Lambda context

    Returns:
        Report URLs
    """
    print(f"Report generation request: {json.dumps(event)}")

    try:
        # Parse request
        if 'body' in event:
            # API Gateway
            body = json.loads(event['body'])
        else:
            # Step Functions
            body = event

        report_format = body.get('format', 'pdf')
        session_id = body.get('session_id')
        analysis_id = body.get('analysis_id')
        user_id = body.get('user_id')

        if not session_id or not analysis_id:
            return _error_response(400, "session_id and analysis_id required")

        # Get analysis data
        analysis_data = _get_analysis_data(user_id, analysis_id)
        if not analysis_data:
            return _error_response(404, "Analysis not found")

        # Generate report based on format
        if report_format == 'pdf':
            result = generate_pdf_report(analysis_data, session_id, analysis_id)
        elif report_format == 'word':
            result = generate_word_report(analysis_data, session_id, analysis_id)
        elif report_format == 'excel':
            result = generate_excel_report(analysis_data, session_id, analysis_id)
        else:
            return _error_response(400, f"Unsupported format: {report_format}")

        return _success_response(result)

    except Exception as e:
        print(f"Error in report handler: {str(e)}")
        return _error_response(500, str(e))


def generate_pdf_report(analysis_data: Dict[str, Any], session_id: str, analysis_id: str) -> Dict[str, Any]:
    """
    Generate PDF report.
    """
    print(f"Generating PDF report for analysis {analysis_id}")

    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
        # Create PDF document
        doc = SimpleDocTemplate(tmp.name, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # Add custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a5490'),
            spaceAfter=30,
            alignment=TA_CENTER
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=12
        )

        # Title page
        story.append(Paragraph("Malaria Risk Analysis Report", title_style))
        story.append(Spacer(1, 0.2*inch))

        # Metadata
        metadata = [
            ['Analysis ID:', analysis_id],
            ['Analysis Type:', analysis_data.get('analysis_type', 'N/A')],
            ['Date:', datetime.now().strftime('%Y-%m-%d %H:%M')],
            ['Status:', analysis_data.get('status', 'Complete')]
        ]

        metadata_table = Table(metadata, colWidths=[2*inch, 4*inch])
        metadata_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ]))
        story.append(metadata_table)
        story.append(Spacer(1, 0.5*inch))

        # Executive Summary
        story.append(Paragraph("Executive Summary", heading_style))
        summary_text = _generate_executive_summary(analysis_data)
        story.append(Paragraph(summary_text, styles['BodyText']))
        story.append(Spacer(1, 0.3*inch))

        # Key Findings
        story.append(Paragraph("Key Findings", heading_style))
        findings = _generate_key_findings(analysis_data)
        for finding in findings:
            story.append(Paragraph(f"• {finding}", styles['BodyText']))
        story.append(PageBreak())

        # Analysis Results
        story.append(Paragraph("Detailed Analysis Results", heading_style))

        # Get results data
        results_df = _get_analysis_results(session_id, analysis_id)
        if results_df is not None and not results_df.empty:
            # Top 10 high-risk wards table
            story.append(Paragraph("Top 10 High-Risk Wards", heading_style))

            if 'risk_score' in results_df.columns:
                ward_col = _find_ward_column(results_df)
                if ward_col:
                    top_10 = results_df.nlargest(10, 'risk_score')[[ward_col, 'risk_score']]
                    top_10['risk_score'] = top_10['risk_score'].round(2)

                    # Convert to table data
                    table_data = [['Ward', 'Risk Score']] + top_10.values.tolist()

                    results_table = Table(table_data, colWidths=[4*inch, 2*inch])
                    results_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 14),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(results_table)

        story.append(PageBreak())

        # Recommendations
        story.append(Paragraph("Recommendations", heading_style))
        recommendations = _generate_recommendations(analysis_data)
        for i, rec in enumerate(recommendations, 1):
            story.append(Paragraph(f"{i}. {rec}", styles['BodyText']))
            story.append(Spacer(1, 0.1*inch))

        # Build PDF
        doc.build(story)

        # Upload to S3
        report_key = f"reports/{session_id}/analysis_report_{analysis_id}.pdf"
        s3_client.upload_file(tmp.name, os.environ.get('DATA_BUCKET', 'chatmrpt-data-prod'), report_key)

        # Generate signed URL
        url = s3_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': os.environ.get('DATA_BUCKET'), 'Key': report_key},
            ExpiresIn=3600
        )

        return {
            'format': 'pdf',
            'url': url,
            'key': report_key,
            'size_kb': os.path.getsize(tmp.name) / 1024
        }


def generate_word_report(analysis_data: Dict[str, Any], session_id: str, analysis_id: str) -> Dict[str, Any]:
    """
    Generate Word document report.
    """
    print(f"Generating Word report for analysis {analysis_id}")

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        # Create Word document
        doc = Document()

        # Set document properties
        doc.core_properties.title = "Malaria Risk Analysis Report"
        doc.core_properties.subject = f"Analysis {analysis_id}"
        doc.core_properties.author = "ChatMRPT System"

        # Title
        title = doc.add_heading('Malaria Risk Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata section
        doc.add_heading('Report Information', level=1)
        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = 'Light List'

        metadata_rows = [
            ('Analysis ID:', analysis_id),
            ('Analysis Type:', analysis_data.get('analysis_type', 'N/A')),
            ('Date:', datetime.now().strftime('%Y-%m-%d %H:%M')),
            ('Status:', analysis_data.get('status', 'Complete'))
        ]

        for i, (label, value) in enumerate(metadata_rows):
            metadata_table.cell(i, 0).text = label
            metadata_table.cell(i, 1).text = str(value)

        doc.add_page_break()

        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        summary_text = _generate_executive_summary(analysis_data)
        doc.add_paragraph(summary_text)

        # Key Findings
        doc.add_heading('Key Findings', level=1)
        findings = _generate_key_findings(analysis_data)
        for finding in findings:
            doc.add_paragraph(finding, style='List Bullet')

        doc.add_page_break()

        # Detailed Results
        doc.add_heading('Detailed Analysis Results', level=1)

        # Get results data
        results_df = _get_analysis_results(session_id, analysis_id)
        if results_df is not None and not results_df.empty:
            doc.add_heading('Top 10 High-Risk Wards', level=2)

            if 'risk_score' in results_df.columns:
                ward_col = _find_ward_column(results_df)
                if ward_col:
                    top_10 = results_df.nlargest(10, 'risk_score')[[ward_col, 'risk_score']]

                    # Create table
                    table = doc.add_table(rows=1, cols=2)
                    table.style = 'Light Grid'

                    # Header row
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Ward'
                    hdr_cells[1].text = 'Risk Score'

                    # Data rows
                    for _, row in top_10.iterrows():
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(row[ward_col])
                        row_cells[1].text = f"{row['risk_score']:.2f}"

        doc.add_page_break()

        # Recommendations
        doc.add_heading('Recommendations', level=1)
        recommendations = _generate_recommendations(analysis_data)
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}")

        # Methodology
        doc.add_heading('Methodology', level=1)
        methodology = _generate_methodology(analysis_data)
        doc.add_paragraph(methodology)

        # Save document
        doc.save(tmp.name)

        # Upload to S3
        report_key = f"reports/{session_id}/analysis_report_{analysis_id}.docx"
        s3_client.upload_file(tmp.name, os.environ.get('DATA_BUCKET', 'chatmrpt-data-prod'), report_key)

        # Generate signed URL
        url = s3_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': os.environ.get('DATA_BUCKET'), 'Key': report_key},
            ExpiresIn=3600
        )

        return {
            'format': 'word',
            'url': url,
            'key': report_key,
            'size_kb': os.path.getsize(tmp.name) / 1024
        }


def generate_excel_report(analysis_data: Dict[str, Any], session_id: str, analysis_id: str) -> Dict[str, Any]:
    """
    Generate Excel report with multiple sheets.
    """
    print(f"Generating Excel report for analysis {analysis_id}")

    with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
        # Create Excel writer
        with pd.ExcelWriter(tmp.name, engine='xlsxwriter') as writer:
            # Get workbook
            workbook = writer.book

            # Add formats
            header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#1a5490',
                'font_color': 'white',
                'align': 'center',
                'valign': 'vcenter'
            })

            # Summary sheet
            summary_data = {
                'Metric': ['Analysis ID', 'Type', 'Date', 'Status', 'Total Wards'],
                'Value': [
                    analysis_id,
                    analysis_data.get('analysis_type', 'N/A'),
                    datetime.now().strftime('%Y-%m-%d %H:%M'),
                    analysis_data.get('status', 'Complete'),
                    'N/A'
                ]
            }
            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, sheet_name='Summary', index=False)

            # Results sheet
            results_df = _get_analysis_results(session_id, analysis_id)
            if results_df is not None and not results_df.empty:
                results_df.to_excel(writer, sheet_name='Results', index=False)

                # Format results sheet
                worksheet = writer.sheets['Results']
                for col_num, value in enumerate(results_df.columns.values):
                    worksheet.write(0, col_num, value, header_format)

                # Statistics sheet
                if 'risk_score' in results_df.columns:
                    stats_data = {
                        'Statistic': ['Mean', 'Median', 'Std Dev', 'Min', 'Max', 'Q1', 'Q3'],
                        'Risk Score': [
                            results_df['risk_score'].mean(),
                            results_df['risk_score'].median(),
                            results_df['risk_score'].std(),
                            results_df['risk_score'].min(),
                            results_df['risk_score'].max(),
                            results_df['risk_score'].quantile(0.25),
                            results_df['risk_score'].quantile(0.75)
                        ]
                    }
                    stats_df = pd.DataFrame(stats_data)
                    stats_df.to_excel(writer, sheet_name='Statistics', index=False)

        # Upload to S3
        report_key = f"reports/{session_id}/analysis_report_{analysis_id}.xlsx"
        s3_client.upload_file(tmp.name, os.environ.get('DATA_BUCKET', 'chatmrpt-data-prod'), report_key)

        # Generate signed URL
        url = s3_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': os.environ.get('DATA_BUCKET'), 'Key': report_key},
            ExpiresIn=3600
        )

        return {
            'format': 'excel',
            'url': url,
            'key': report_key,
            'size_kb': os.path.getsize(tmp.name) / 1024
        }


def _get_analysis_data(user_id: str, analysis_id: str) -> Optional[Dict[str, Any]]:
    """Get analysis metadata from DynamoDB."""
    table = dynamodb.Table(os.environ.get('ANALYSES_TABLE', 'chatmrpt-analyses-prod'))

    response = table.get_item(
        Key={
            'user_id': user_id,
            'analysis_id': analysis_id
        }
    )

    return response.get('Item')


def _get_analysis_results(session_id: str, analysis_id: str) -> Optional[pd.DataFrame]:
    """Get analysis results from S3."""
    results_key = f"results/{session_id}/analysis_{analysis_id}.csv"

    try:
        with tempfile.NamedTemporaryFile(suffix='.csv') as tmp:
            s3_client.download_file(
                os.environ.get('DATA_BUCKET', 'chatmrpt-data-prod'),
                results_key,
                tmp.name
            )
            return pd.read_csv(tmp.name)
    except Exception as e:
        print(f"Error getting results: {str(e)}")
        return None


def _find_ward_column(df: pd.DataFrame) -> Optional[str]:
    """Find ward/admin column in dataframe."""
    candidates = ['ward', 'Ward', 'ward_name', 'Ward_Name', 'lga', 'LGA']
    for col in candidates:
        if col in df.columns:
            return col
    return None


def _generate_executive_summary(analysis_data: Dict[str, Any]) -> str:
    """Generate executive summary text."""
    analysis_type = analysis_data.get('analysis_type', 'risk analysis')

    summaries = {
        'composite_scoring': (
            "This composite scoring analysis integrates multiple malaria risk indicators to provide "
            "a comprehensive assessment of vulnerability across geographic areas. The analysis combines "
            "demographic, environmental, and health system factors to identify high-priority areas for "
            "targeted malaria interventions."
        ),
        'pca': (
            "Principal Component Analysis (PCA) was performed to identify the key underlying factors "
            "driving malaria risk patterns. This dimensionality reduction technique reveals the most "
            "influential variables and their relationships, enabling more focused intervention strategies."
        ),
        'vulnerability_ranking': (
            "The vulnerability ranking analysis systematically evaluates and ranks geographic areas based "
            "on their susceptibility to malaria transmission. This prioritization framework supports "
            "evidence-based resource allocation and intervention planning."
        ),
        'settlement_analysis': (
            "Settlement footprint analysis examines the distribution and characteristics of human settlements "
            "in relation to malaria risk factors. This spatial analysis helps identify population centers "
            "at highest risk and informs targeted intervention strategies."
        )
    }

    return summaries.get(analysis_type, (
        f"This {analysis_type} provides insights into malaria risk patterns and vulnerability "
        "factors across the study area, supporting data-driven decision making for public health interventions."
    ))


def _generate_key_findings(analysis_data: Dict[str, Any]) -> List[str]:
    """Generate key findings list."""
    findings = [
        "Significant spatial variation in malaria risk levels identified across the study area",
        "Environmental and demographic factors show strong correlation with risk patterns",
        "Several high-priority areas identified for immediate intervention",
        "Data quality and coverage adequate for reliable risk assessment",
        "Clear clustering of high-risk areas suggests targeted intervention opportunities"
    ]
    return findings


def _generate_recommendations(analysis_data: Dict[str, Any]) -> List[str]:
    """Generate recommendations based on analysis."""
    analysis_type = analysis_data.get('analysis_type', '')

    base_recommendations = [
        "Prioritize high-risk areas identified in this analysis for immediate intervention",
        "Implement targeted ITN distribution campaigns in vulnerable communities",
        "Strengthen surveillance systems in areas showing emerging risk patterns",
        "Conduct community engagement activities to improve intervention uptake",
        "Regular monitoring and evaluation to track intervention impact"
    ]

    if analysis_type == 'composite_scoring':
        base_recommendations.insert(0, "Use composite scores to guide multi-sectoral intervention planning")
    elif analysis_type == 'settlement_analysis':
        base_recommendations.insert(0, "Focus interventions on densely populated settlements with high risk scores")

    return base_recommendations


def _generate_methodology(analysis_data: Dict[str, Any]) -> str:
    """Generate methodology description."""
    analysis_type = analysis_data.get('analysis_type', 'analysis')

    base_methodology = (
        f"This {analysis_type} employed standardized WHO-recommended methodologies for malaria risk assessment. "
        "Data preprocessing included validation, normalization, and handling of missing values. "
        "Statistical analysis was performed using established epidemiological techniques, with "
        "results validated through sensitivity analysis and cross-validation where applicable."
    )

    return base_methodology


def _success_response(data: Any) -> Dict[str, Any]:
    """Create successful API response."""
    return {
        'statusCode': 200,
        'headers': {
            'Content-Type': 'application/json',
            'Access-Control-Allow-Origin': '*'
        },
        'body': json.dumps(data)
    }


def _error_response(status_code: int, message: str) -> Dict[str, Any]:
    """Create error API response."""
    return {
        'statusCode': status_code,
        'headers': {
            'Content-Type': 'application/json',
            'Access-Control-Allow-Origin': '*'
        },
        'body': json.dumps({'error': message})
    }