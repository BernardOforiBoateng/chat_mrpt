"""
Extract comments from Word document
"""

import docx
from docx import Document
import zipfile
import xml.etree.ElementTree as ET
from lxml import etree
import json

def extract_comments_from_docx(filepath):
    """Extract comments from a Word document"""
    
    # Load the document
    doc = Document(filepath)
    
    # Also need to access the raw XML for comments
    with zipfile.ZipFile(filepath, 'r') as docx_zip:
        # Check if comments exist
        if 'word/comments.xml' in docx_zip.namelist():
            comments_xml = docx_zip.read('word/comments.xml')
            comments_tree = etree.fromstring(comments_xml)
            
            # Define namespaces
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            # Extract comments
            comments = []
            for comment in comments_tree.xpath('//w:comment', namespaces=namespaces):
                comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')
                date = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
                
                # Extract comment text
                comment_text = ''
                for para in comment.xpath('.//w:p', namespaces=namespaces):
                    for text_elem in para.xpath('.//w:t', namespaces=namespaces):
                        comment_text += text_elem.text if text_elem.text else ''
                    comment_text += '\n'
                
                comments.append({
                    'id': comment_id,
                    'author': author,
                    'date': date,
                    'text': comment_text.strip()
                })
            
            return comments
        else:
            return []

def extract_document_content(filepath):
    """Extract main document content"""
    doc = Document(filepath)
    
    content = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            content.append({
                'paragraph': i + 1,
                'text': para.text.strip()
            })
    
    return content

def main():
    filepath = '/mnt/c/Users/bbofo/OneDrive/Desktop/ChatMRPT/tpr_analysis/TPR_Final_Deliverables/TPR_COMPREHENSIVE_FINDINGS_REPORT_WITH_COMMENTS.docx'
    
    print("="*80)
    print("EXTRACTING COMMENTS FROM TPR REPORT")
    print("="*80)
    
    # Extract comments
    comments = extract_comments_from_docx(filepath)
    
    if comments:
        print(f"\nFound {len(comments)} comments in the document:\n")
        print("-"*80)
        
        for i, comment in enumerate(comments, 1):
            print(f"\nCOMMENT #{i}")
            print(f"Author: {comment['author']}")
            print(f"Date: {comment['date'][:10] if comment['date'] else 'N/A'}")
            print(f"Comment Text:")
            print("-"*40)
            print(comment['text'])
            print("-"*80)
    else:
        print("\nNo comments found in the document.")
    
    # Also extract key sections of the document
    print("\n" + "="*80)
    print("EXTRACTING KEY DOCUMENT SECTIONS")
    print("="*80)
    
    content = extract_document_content(filepath)
    
    # Find headings and key sections
    headings = []
    for item in content:
        text = item['text']
        # Identify potential headings (usually in all caps or numbered sections)
        if (text.isupper() and len(text) < 100) or text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            headings.append(text)
    
    print("\nDocument Structure (Main Headings):")
    for heading in headings[:20]:  # Show first 20 headings
        print(f"  - {heading}")
    
    # Save results to file
    output = {
        'comments': comments,
        'document_sections': headings[:50],
        'total_paragraphs': len(content)
    }
    
    with open('/mnt/c/Users/bbofo/OneDrive/Desktop/ChatMRPT/tpr_analysis/extracted_comments.json', 'w') as f:
        json.dump(output, f, indent=2)
    
    print(f"\nResults saved to extracted_comments.json")
    print(f"Total paragraphs in document: {len(content)}")
    print(f"Total comments found: {len(comments)}")

if __name__ == "__main__":
    main()